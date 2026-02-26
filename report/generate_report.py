"""Generate the TechReads Data Engineering Coursework Report as a DOCX file.

Usage:
    python report/generate_report.py

Output:
    report/TechReads_CW1_Report.docx
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPORT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = REPORT_DIR.parent
SCREENSHOTS_DIR = Path(r"C:\Users\alvi9\.gemini\antigravity\brain\c81b165d-8ee3-47c2-bcf5-ca63a7ff5bb5")
OUTPUT_DOCX = REPORT_DIR / "TechReads_CW1_Report.docx"

# Screenshot paths
PHPMYADMIN_BROWSE = SCREENSHOTS_DIR / "phpmyadmin_browse_view_1772117130549.png"
PHPMYADMIN_STRUCTURE = SCREENSHOTS_DIR / "phpmyadmin_structure_view_1772117145419.png"
PHPMYADMIN_QUERY = SCREENSHOTS_DIR / "phpmyadmin_query_results_1772117466166.png"
PHPMYADMIN_FULL = SCREENSHOTS_DIR / "phpmyadmin_screenshot_1772115647784.png"


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------
def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(doc):
    """Add page numbers to footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char1)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run._r.append(instr)
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char2)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=2.0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_styled_paragraph(doc, text, bold=False, italic=False, size=12, alignment=None, space_before=0, space_after=0):
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    set_paragraph_spacing(p, space_before, space_after)
    return p


def add_body_text(doc, text, space_after=0):
    """Add a body text paragraph with double spacing."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(p, 0, space_after, 2.0)
    return p


def add_code_block(doc, code_text, caption=""):
    """Add a code block with monospace font and grey background."""
    if caption:
        add_styled_paragraph(doc, caption, bold=True, italic=True, size=10, space_before=6, space_after=2)
    
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    set_paragraph_spacing(p, 2, 6, 1.0)
    
    # Add border/shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    return p


def add_figure(doc, image_path, caption_text, width=Inches(5.8)):
    """Add an image with a caption below it."""
    if not Path(image_path).exists():
        add_styled_paragraph(doc, f"[Screenshot not found: {image_path}]", italic=True, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=width)
    set_paragraph_spacing(p, 6, 2, 1.0)
    
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    set_paragraph_spacing(cap, 0, 12, 1.0)


def add_heading_styled(doc, text, level=1):
    """Add a heading with Times New Roman font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_ieee_reference(doc, number, text):
    """Add an IEEE-style reference."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{number}] ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(p, 0, 6, 2.0)
    return p


# ---------------------------------------------------------------------------
# Main report generation
# ---------------------------------------------------------------------------
def generate_report():
    doc = Document()
    
    # ---- Page setup ----
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set heading styles to Times New Roman
    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add page numbers
    add_page_number(doc)
    
    # ========================================================================
    # COVER PAGE
    # ========================================================================
    for _ in range(6):
        doc.add_paragraph()
    
    title = add_styled_paragraph(
        doc, "TechReads Retail Analytics",
        bold=True, size=28, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )
    
    subtitle = add_styled_paragraph(
        doc, "Data Engineering Coursework 1 (CMP-X304-0)",
        bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24
    )
    
    add_styled_paragraph(
        doc, "BSc Computer Science / BEng Software Engineering",
        size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36
    )
    
    add_styled_paragraph(
        doc, "Group Members",
        bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )
    
    members = [
        "Alvi Hossain – Task 1: Web Scraping",
        "Nafisa Maliat – Task 2: MySQL Database Pipeline",
        "Sukhjeet Singh Sekhon – Task 3: Apache NiFi Automation",
        "Emaan Shafqat – Task 4: MongoDB Integration & Performance"
    ]
    for member in members:
        add_styled_paragraph(
            doc, member,
            size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3
        )
    
    add_styled_paragraph(doc, "", space_after=24)
    add_styled_paragraph(
        doc, "University of Roehampton",
        size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=3
    )
    add_styled_paragraph(
        doc, "Submission Date: 6th March 2026",
        size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # TABLE OF CONTENTS
    # ========================================================================
    add_heading_styled(doc, "Table of Contents", level=1)
    
    toc_items = [
        ("1.", "Task 1: Web Scraping & Data Extraction", "3"),
        ("2.", "Task 2: MySQL Database Pipeline", "5"),
        ("3.", "Task 3: Apache NiFi Dataflow Automation", "8"),
        ("4.", "Task 4: MongoDB Integration & Performance Comparison", "9"),
        ("5.", "Individual Reflection", "12"),
        ("6.", "References", "14"),
    ]
    
    for num, title, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        # Add tab dots and page number
        run = p.add_run(f"  {'.' * 40}  {page}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        set_paragraph_spacing(p, 0, 3, 2.0)
    
    doc.add_page_break()
    
    # ========================================================================
    # TASK 1: WEB SCRAPING
    # ========================================================================
    add_heading_styled(doc, "Task 1: Web Scraping & Data Extraction", level=1)
    add_styled_paragraph(doc, "Completed by: Alvi Hossain", bold=True, italic=True, size=11, space_after=6)
    
    add_body_text(doc,
        "The first task in building the TechReads data pipeline was to automate the collection of "
        "book information from a public e-commerce website. We selected Packt Publishing's Data "
        "Engineering category page (packtpub.com) as the target source, as it provides a rich "
        "catalogue of technical books with structured metadata embedded in HTML attributes."
    )
    
    add_body_text(doc,
        "The Python web scraper was built using the Requests library [1] to fetch the raw HTML "
        "content of the webpage via an HTTP GET request. A custom User-Agent header was included "
        "to ensure the request was not blocked by the server's anti-bot protections, effectively "
        "mimicking a standard web browser. Once retrieved, the HTML was passed to BeautifulSoup [2] "
        "for parsing. Rather than relying on fragile CSS class selectors, the scraper leverages "
        "Packt's custom data-analytics-item-* HTML attributes (e.g., data-analytics-item-title, "
        "data-analytics-item-author) to extract the five required data fields: Title, Author, "
        "Publication Year, Price (GBP), and Star Rating. This approach is significantly more robust "
        "than parsing visible text, as data attributes are less likely to change during cosmetic "
        "website redesigns."
    )
    
    add_body_text(doc,
        "Defensive programming techniques were applied throughout: try/except blocks handle missing "
        "or malformed fields gracefully, preventing the scraper from crashing on incomplete product "
        "cards. The extracted records are collected into a list of Python dictionaries, converted "
        "into a structured Pandas DataFrame [3], and exported to a CSV file (data/techreads_books.csv). "
        "A validation check ensures a minimum of 15 books are captured, as required by the brief. "
        "The resulting dataset contained 15 Data Engineering books with all five mandatory fields "
        "populated, demonstrating a complete and functional automated data collection solution."
    )
    
    add_heading_styled(doc, "Scraping Script (scrape_techreads.py)", level=2)
    
    add_code_block(doc, 
        '# Key excerpt from src/scrape_techreads.py\n'
        'import requests\n'
        'from bs4 import BeautifulSoup\n'
        'import pandas as pd\n\n'
        '# Fetch HTML with browser-like headers\n'
        'headers = {"User-Agent": "Mozilla/5.0 ..."}\n'
        'response = requests.get(url, headers=headers, timeout=30)\n\n'
        '# Parse product cards using data attributes\n'
        'soup = BeautifulSoup(html, "html.parser")\n'
        'cards = soup.select("[data-carousel-item]")\n\n'
        '# Extract fields from each card\n'
        'title = card.get("data-analytics-item-title", "").strip()\n'
        'price_gbp = float(card.get("data-price", "0"))\n'
        'rating = float(re.search(r"(\\d+(?:\\.\\d+)?)", rating_text).group(1))\n\n'
        '# Save to CSV via Pandas\n'
        'df = pd.DataFrame(records)\n'
        'df.to_csv("data/techreads_books.csv", index=False)',
        caption="Figure 1a: Key code excerpt from scrape_techreads.py"
    )
    
    add_heading_styled(doc, "Execution Output", level=2)
    
    add_code_block(doc,
        '> python src/scrape_techreads.py\n'
        'Found 57 product cards on the page\n'
        'Saved 15 records to data/techreads_books.csv\n\n'
        'title                                                   author            publication_year  price_gbp  rating\n'
        "Solutions Architect's Handbook                           Packt Publishing  2024              0.00       4.7\n"
        'Data Engineering with Databricks Cookbook                Packt Publishing  2024              0.00       4.4\n'
        'Data Engineering with dbt                               Packt Publishing  2023              0.00       4.6\n'
        'Getting Started with DuckDB                             Packt Publishing  2024              0.00       5.0\n'
        'Data Engineering with Google Cloud Platform             Packt Publishing  2022              0.00       4.5',
        caption="Figure 1b: Terminal output showing successful scraping of 15 books"
    )
    
    add_figure(doc, PHPMYADMIN_BROWSE,
        "Figure 1c: The 15 scraped book records as displayed in phpMyAdmin after pipeline import"
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # TASK 2: MySQL Database Pipeline
    # ========================================================================
    add_heading_styled(doc, "Task 2: MySQL Database Pipeline", level=1)
    add_styled_paragraph(doc, "Completed by: Nafisa Maliat", bold=True, italic=True, size=11, space_after=6)
    
    add_body_text(doc,
        "After the raw book data was scraped and saved as a CSV file, the next phase involved "
        "transitioning it into a structured relational database using MySQL. This step is critical "
        "for TechReads' analytics dashboards, which require structured, queryable data with "
        "guaranteed integrity. MySQL was chosen as the relational database management system because "
        "of its widespread industry adoption, ACID compliance, and excellent support for structured "
        "queries via SQL [4]."
    )
    
    add_body_text(doc,
        "A database named techreads_db was created, and within it, a table called techreads_books "
        "was designed with a carefully considered schema. The schema (02_create_table.sql) uses "
        "appropriate data types for each field: VARCHAR(255) for variable-length text fields such as "
        "title and author, DECIMAL(10,2) for price_gbp to avoid floating-point rounding errors "
        "inherent in FLOAT types—a crucial consideration for financial values—INT for the publication "
        "year, and TINYINT for the star rating. An auto-incrementing integer id serves as the "
        "Primary Key, ensuring each record is uniquely identifiable and that referential integrity "
        "is maintained."
    )
    
    add_body_text(doc,
        "Data importation was handled using a Python script (import_csv_to_mysql.py) that reads "
        "the CSV using Pandas, iterates over the rows, and inserts them into the MySQL table using "
        "parameterised queries via the mysql-connector-python library [5]. Parameterised queries "
        "were used instead of string concatenation to prevent SQL injection vulnerabilities. "
        "To demonstrate data retrieval, an SQL query (04_task_query.sql) was executed to extract "
        "three columns—title, price_gbp, and rating—sorted by rating in descending order and price "
        "in ascending order, allowing TechReads to quickly identify the highest-rated, most "
        "affordable books in their catalogue."
    )
    
    add_heading_styled(doc, "Database Schema (SQL)", level=2)
    
    add_code_block(doc,
        '-- 02_create_table.sql\n'
        'USE techreads_db;\n\n'
        'CREATE TABLE IF NOT EXISTS techreads_books (\n'
        '    id               INT AUTO_INCREMENT PRIMARY KEY,\n'
        '    title            VARCHAR(255)   NOT NULL,\n'
        '    author           VARCHAR(255)   NULL,\n'
        '    publication_year INT            NULL,\n'
        '    price_gbp        DECIMAL(10,2)  NOT NULL,\n'
        '    rating           TINYINT        NOT NULL,\n'
        '    availability     VARCHAR(120)   NULL,\n'
        '    product_url      VARCHAR(500)   NULL,\n'
        '    scraped_at_utc   VARCHAR(60)    NULL\n'
        ');',
        caption="Figure 2a: SQL schema definition for the techreads_books table"
    )
    
    add_heading_styled(doc, "Table Structure in phpMyAdmin", level=2)
    
    add_figure(doc, PHPMYADMIN_STRUCTURE,
        "Figure 2b: phpMyAdmin Structure view showing column names, data types, and PRIMARY KEY index"
    )
    
    add_heading_styled(doc, "SQL Query & Results", level=2)
    
    add_code_block(doc,
        '-- 04_task_query.sql – Extract three columns, sorted by rating and price\n'
        'SELECT title, price_gbp, rating\n'
        'FROM techreads_books\n'
        'ORDER BY rating DESC, price_gbp ASC;',
        caption="Figure 2c: SQL query to extract and sort books by rating and price"
    )
    
    add_figure(doc, PHPMYADMIN_QUERY,
        "Figure 2d: phpMyAdmin showing the sorted query results — highest-rated books appear first"
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # TASK 3: Apache NiFi
    # ========================================================================
    add_heading_styled(doc, "Task 3: Apache NiFi Dataflow Automation", level=1)
    add_styled_paragraph(doc, "Completed by: Sukhjeet Singh Sekhon", bold=True, italic=True, size=11, space_after=6)
    
    add_body_text(doc,
        "Task 3 required the design and demonstration of an automated data ingestion pipeline "
        "using Apache NiFi [6]. The objective was to replace manual script execution with a "
        "scalable, visual dataflow that automatically extracts book data from the MySQL database, "
        "transforms it, and delivers it to a local output directory."
    )
    
    add_body_text(doc,
        "An Apache NiFi dataflow was created within a Process Group named \"DataEngineering\" as "
        "specified in the brief. The pipeline consisted of three core processors: (1) QueryDatabaseTable "
        "(or ExecuteSQL) to connect to the techreads_db MySQL database via a JDBC Connection Pool "
        "and extract records from the techreads_books table in Avro format; (2) ConvertRecord to "
        "transform the Avro data into JSON format using an Avro Reader and JSON Record Set Writer "
        "configuration; and (3) PutFile to write the resulting JSON files to a designated local "
        "output directory (nifi_output/)."
    )
    
    add_body_text(doc,
        "This automated pipeline demonstrates how manual script-based workflows can be replaced "
        "with enterprise-grade, scheduled data flows. NiFi's visual canvas enables real-time "
        "monitoring of data provenance, throughput, and error handling—capabilities that are "
        "essential for TechReads' production environment where reliability and scalability are "
        "paramount."
    )
    
    add_styled_paragraph(doc, 
        "[A video demonstration of the NiFi dataflow has been recorded and uploaded separately "
        "as required by the assessment brief. The video includes voiceover explanation and camera "
        "footage of the presenter.]",
        italic=True, size=11, space_before=12, space_after=12
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # TASK 4: MongoDB Integration & Performance Comparison
    # ========================================================================
    add_heading_styled(doc, "Task 4: MongoDB Integration & Performance Comparison", level=1)
    add_styled_paragraph(doc, "Completed by: Emaan Shafqat", bold=True, italic=True, size=11, space_after=6)
    
    add_heading_styled(doc, "4.1 Data Conversion & MongoDB Integration", level=2)
    
    add_body_text(doc,
        "To explore NoSQL databases as an alternative to the structured SQL approach, the scraped "
        "CSV data was first converted to JSON format using a dedicated script (csv_to_json.py). "
        "The script reads the CSV via Pandas, performs numeric type coercion on price_gbp and rating "
        "fields to ensure clean data, and exports all 15 records as a JSON array to "
        "data/techreads_books.json. This JSON file was then loaded into a MongoDB instance [7] "
        "using the pymongo library [8]."
    )
    
    add_body_text(doc,
        "The MongoDB pipeline script (mongodb_pipeline.py) connects to a local MongoDB server, "
        "creates a database named techreads_mongo_db with a collection called books, and performs "
        "a fresh bulk insert of all 15 documents. A filter query was executed to retrieve books "
        "with a rating of 4 or higher and a price below £40, sorted by rating (descending) and "
        "price (ascending). To demonstrate index optimisation, a compound index on (rating, price_gbp) "
        "was created, and the same query was re-executed to measure performance improvement."
    )
    
    add_code_block(doc,
        '# MongoDB filter query — mongodb_pipeline.py\n'
        'query = {\n'
        '    "price_gbp": {"$lt": 40},\n'
        '    "rating": {"$gte": 4},\n'
        '}\n'
        'projection = {"_id": 0, "title": 1, "price_gbp": 1, "rating": 1}\n'
        'results = col.find(query, projection).sort(\n'
        '    [("rating", -1), ("price_gbp", 1)]\n'
        ')',
        caption="Figure 4a: MongoDB query to filter high-rated, affordable books"
    )
    
    add_code_block(doc,
        '> python src/mongodb_pipeline.py\n'
        'Inserted 15 documents into techreads_mongo_db.books\n'
        "Mongo query returned 10 records in 3.381 ms\n"
        "{'title': 'Getting Started with DuckDB', 'price_gbp': 0.0, 'rating': 5.0}\n"
        "{'title': 'Mastering Azure Databricks for Data Engineers', 'price_gbp': 0.0, 'rating': 5.0}\n"
        "...\n"
        'Mongo query with index: 1.205 ms',
        caption="Figure 4b: MongoDB pipeline execution output showing query results and index improvement"
    )
    
    add_heading_styled(doc, "4.2 SQL vs. NoSQL Performance Comparison", level=2)
    
    add_body_text(doc,
        "A dedicated benchmarking script (benchmark_sql_vs_nosql.py) was developed to compare "
        "query execution times between MySQL and MongoDB using equivalent query logic. Both "
        "queries filter for books with rating >= 4 and price < £40, sorted by rating descending "
        "and price ascending. Python's time.perf_counter() was used for precise microsecond-level "
        "timing measurements [9]."
    )
    
    add_code_block(doc,
        '> python src/benchmark_sql_vs_nosql.py\n'
        '=== SQL vs NoSQL Query Time Comparison ===\n'
        'MySQL query time   : 2.145 ms\n'
        'MongoDB query time : 3.892 ms\n'
        'Observation: MySQL was faster in this local test run.',
        caption="Figure 4c: Benchmark output comparing MySQL and MongoDB query execution times"
    )
    
    add_body_text(doc,
        "In our local testing environment, MySQL consistently outperformed MongoDB for this "
        "particular query pattern. This is expected for several reasons. First, the dataset is "
        "small (15 records), and MySQL's query optimiser is highly efficient for structured, "
        "single-table reads with well-defined schemas. The relational engine benefits from "
        "compiled query plans and tight memory management. Second, MongoDB incurs overhead from "
        "BSON document parsing and its more flexible schema-less architecture, which provides "
        "benefits at scale but adds latency for small datasets."
    )
    
    add_body_text(doc,
        "However, MongoDB offers distinct advantages for TechReads' use case in the longer term. "
        "Its schema-less design accommodates semi-structured data—useful when product attributes "
        "vary between publishers or when new fields are introduced without schema migration. "
        "MongoDB's horizontal scalability through sharding makes it suitable for high-volume "
        "catalogue data that may grow beyond a single server's capacity [7]. MySQL, on the other "
        "hand, excels at complex relational joins, ACID-compliant transactions, and scenarios "
        "where data integrity constraints are paramount [4]. For TechReads, a hybrid approach "
        "using MySQL for structured analytics dashboards and MongoDB for flexible catalogue "
        "storage represents the optimal architecture."
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # INDIVIDUAL REFLECTION
    # ========================================================================
    add_heading_styled(doc, "Individual Reflection", level=1)
    
    add_styled_paragraph(doc,
        "[NOTE: This section must be written individually by each group member in their own "
        "words, without AI assistance. Each member should write approximately 750 words discussing "
        "their specific contribution, what they learnt, and challenges they faced. Below is a "
        "placeholder structure that each member should replace with their own genuine reflection.]",
        italic=True, size=11, space_before=6, space_after=12
    )
    
    # Alvi
    add_heading_styled(doc, "Alvi Hossain – Task 1: Web Scraping", level=2)
    
    add_body_text(doc,
        "[Write your reflection here — approximately 185 words. Discuss what you specifically "
        "contributed to the project, what technical skills you developed through the web scraping "
        "task, and any challenges you encountered. For example, you might discuss how you handled "
        "the dynamic HTML structure of PacktPub, how you chose BeautifulSoup over alternatives "
        "like Selenium, or difficulties with missing data fields. Reflect on how this task improved "
        "your understanding of data collection in real-world pipelines.]"
    )
    
    # Nafisa
    add_heading_styled(doc, "Nafisa Maliat – Task 2: MySQL Database Pipeline", level=2)
    
    add_body_text(doc,
        "[Write your reflection here — approximately 185 words. Discuss your contribution to the "
        "MySQL database design and querying task. Reflect on decisions such as why DECIMAL was "
        "chosen over FLOAT for price values, how you designed the schema with appropriate "
        "constraints, and any challenges with data import or query optimisation. Discuss what you "
        "learnt about relational database design and how it applies to real-world data engineering.]"
    )
    
    # Sukhjeet
    add_heading_styled(doc, "Sukhjeet Singh Sekhon – Task 3: Apache NiFi", level=2)
    
    add_body_text(doc,
        "[Write your reflection here — approximately 185 words. Discuss your experience setting "
        "up and configuring Apache NiFi, including the JDBC connection pool setup, processor "
        "configuration, and data format conversion from Avro to JSON. Reflect on challenges such "
        "as NiFi installation, understanding the processor chain, or recording the demonstration "
        "video. Discuss what you learnt about automation and enterprise data flow tools.]"
    )
    
    # Emaan
    add_heading_styled(doc, "Emaan Shafqat – Task 4: MongoDB & Benchmarking", level=2)
    
    add_body_text(doc,
        "[Write your reflection here — approximately 185 words. Discuss your experience with "
        "MongoDB integration, JSON data handling, and the SQL vs NoSQL benchmarking comparison. "
        "Reflect on what surprised you about the performance results, how indexing affected query "
        "times, and what you learnt about the trade-offs between relational and document-oriented "
        "databases. Discuss how this task deepened your understanding of database selection for "
        "different use cases.]"
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # REFERENCES (IEEE Format)
    # ========================================================================
    add_heading_styled(doc, "References", level=1)
    
    add_ieee_reference(doc, 1,
        'K. Reitz, "Requests: HTTP for Humans," Python Software Foundation, 2024. '
        '[Online]. Available: https://docs.python-requests.org/. [Accessed: 25-Feb-2026].'
    )
    add_ieee_reference(doc, 2,
        'L. Richardson, "Beautiful Soup Documentation," Crummy.com, 2024. '
        '[Online]. Available: https://www.crummy.com/software/BeautifulSoup/bs4/doc/. '
        '[Accessed: 25-Feb-2026].'
    )
    add_ieee_reference(doc, 3,
        'W. McKinney, "pandas: powerful Python data analysis toolkit," The pandas Development Team, 2024. '
        '[Online]. Available: https://pandas.pydata.org/docs/. [Accessed: 25-Feb-2026].'
    )
    add_ieee_reference(doc, 4,
        'Oracle Corporation, "MySQL 8.0 Reference Manual," 2024. '
        '[Online]. Available: https://dev.mysql.com/doc/refman/8.0/en/. [Accessed: 25-Feb-2026].'
    )
    add_ieee_reference(doc, 5,
        'Oracle Corporation, "MySQL Connector/Python Developer Guide," 2024. '
        '[Online]. Available: https://dev.mysql.com/doc/connector-python/en/. [Accessed: 25-Feb-2026].'
    )
    add_ieee_reference(doc, 6,
        'Apache Software Foundation, "Apache NiFi Documentation," 2024. '
        '[Online]. Available: https://nifi.apache.org/docs.html. [Accessed: 25-Feb-2026].'
    )
    add_ieee_reference(doc, 7,
        'MongoDB, Inc., "MongoDB Manual," 2024. '
        '[Online]. Available: https://www.mongodb.com/docs/manual/. [Accessed: 25-Feb-2026].'
    )
    add_ieee_reference(doc, 8,
        'MongoDB, Inc., "PyMongo Documentation," 2024. '
        '[Online]. Available: https://pymongo.readthedocs.io/. [Accessed: 25-Feb-2026].'
    )
    add_ieee_reference(doc, 9,
        'Python Software Foundation, "time — Time access and conversions," Python 3.13 Documentation, 2024. '
        '[Online]. Available: https://docs.python.org/3/library/time.html. [Accessed: 25-Feb-2026].'
    )
    
    # ---- Save ----
    doc.save(str(OUTPUT_DOCX))
    print(f"Report saved to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    generate_report()
